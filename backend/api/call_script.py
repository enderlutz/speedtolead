"""
Call script API. Single-row table — admin edits one master script at
Settings → Call Script. The VA's sticky panel on Lead Detail pulls
the raw template + the active lead, substitutes variables + conditional
blocks client-side, and renders the result.

The initial seed content is the script Olga uses today, marked up with
{{var}} placeholders + {{#if X}}/{{/if}} branches matching the
"if they put X / if they didn't" forks in the original script.
"""
from __future__ import annotations
import logging
import re
import uuid
from datetime import datetime, timezone
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File
from pydantic import BaseModel
from sqlalchemy import func

from database import get_db, CallScript
from api.auth import require_admin, get_current_user

router = APIRouter()
logger = logging.getLogger(__name__)

DEFAULT_SCRIPT_ID = "default"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


# Initial template seeded into the DB on first read. Admin can edit/replace
# from the Settings page; this string is only the fallback if the row is
# missing. Variables: customer_name, customer_first_name, your_name, brand,
# address, zip_code, fence_height, fence_age, previously_stained,
# fence_sides, linear_feet, tier_essential, tier_signature, tier_legacy.
# Conditionals: has_address, has_fence_height, has_fence_age,
# has_previously_stained, has_fence_sides, fence_brand_new, fence_older_than_6mo.
INITIAL_SCRIPT = """## Stage 1 — Greeting

Hi, is this {{customer_name}}? Hi {{customer_first_name}}! This is {{your_name}} calling from {{brand}} — we just received your information. How are you doing?

> **Customer:** I'm doing great, how are you?
>
> **You:** I'm doing great
>
> **Customer:** I'm sorry, who is this?
>
> **You:** This is {{your_name}} with {{brand}}

## Stage 2 — Address Confirmation

I just want to call to make sure everything we have on file is correct so that I can send your estimate.

{{#if has_address}}
I have your address as **{{address}}, {{zip_code}}**. Is that correct?

*(update it if anything is incorrect)*
{{/if}}{{#unless has_address}}
What's the address where you'd like the fence staining done? And what's the ZIP code there? I want to make sure we get an accurate measurement using Google Earth.

Just to confirm — that's **[FULL ADDRESS, ZIP]**?

*(if they don't want to give you an address then ask them for the linear feet so we can give them a rough estimate)*
{{/unless}}

## Stage 3 — Fence Details

### Fence Height

{{#if has_fence_height}}
I just wanted to confirm that your fence is **{{fence_height}}** tall?

*(ex: six, six and a half, seven, seven and a half)*
{{/if}}{{#unless has_fence_height}}
*(confirm it for yourself later and don't talk to them about it)*
{{/unless}}

### Fence Age

{{#if has_fence_age}}
I see here that your fence is between **{{fence_age}}** old?
{{/if}}{{#unless has_fence_age}}
About how old would you say the fence is?
{{/unless}}

### Previously Stained

{{#if has_previously_stained}}
And it is **{{previously_stained}}** previously stained, right?
{{/if}}{{#unless has_previously_stained}}
*(confirm it for yourself later and don't talk to them about it)*
{{/unless}}

### Fence Sides

Okay and one last thing! — which sides of the fence do you want stained? Which sides?

**8 possible sides:**
- Inside front *(usually contains a gate)*
- Inside left
- Inside back
- Inside right
- Outside front *(usually contains a gate)*
- Outside left
- Outside back
- Outside right

*If they pick all 4 inside:* "OK I'll quote you for all the insides"

Once they tell you, **confirm back to them**:

> Okay so to confirm you want the *(name the sides)*?

{{#if has_fence_sides}}
*(They previously selected: {{fence_sides}})*
{{/if}}

## Stage 4 — Package Walkthrough

{{#if fence_brand_new}}
OK perfect, since your fence is brand new, it will not require a cleaning before we stain it. I will make sure to take that charge off the cleaning fee since it won't be needed.
{{/if}}{{#if fence_older_than_6mo}}
OK since your fence is over 6 months old, it will need to be cleaned before we stain it. We use a biodegradable chemical wash that is safe for the plants and pets. It removes the grey color, sprinkler stains, and any mold — and we also pressure wash if it's needed.

This will make sure the stain penetrates the wood properly.
{{/if}}

## Stage 5 — Explain the Packages

Once we clean and stain it we will apply the stain based on your chosen package which we will have on the estimate. Would you like me to explain the packages or would you prefer to review them when I send the estimate?

**The first package is the Essential** — it's a clear stain that protects the fence and lasts one to three years.

**The second package is the Signature Finish** — it's our most popular package. The colors are rich and have more of a natural look that lasts three to six years.

**The third package is the Legacy Premium** — a solid and bold color that gives you the maximum protection and typically lasts three to eight years.

Okay, and just to let you know, every package comes with **2 coats of stain** and we also **back-brush** the stain.

## Stage 6 — Set Estimate Expectations

Perfect, once I get off the call I'm going to measure your fence on Google Earth and then I'll text you your estimate here in a few minutes. Do you have any questions before I send you the estimate?

## Stage 7 — Goodbye

Okay, thank you, have a great day!

---

**If you don't know the answer:**

> "That's a great question for our project manager. I'll have him reach out to you after this call!"
"""


class CallScriptBody(BaseModel):
    content: str


class CallScriptCreate(BaseModel):
    name: str
    content: str = ""


class CallScriptUpdate(BaseModel):
    name: str | None = None
    content: str | None = None


def _ensure_seed(db) -> CallScript:
    """Lazy seed — first read on a fresh DB inserts the initial script as the
    library's first named entry. Subsequent reads return whatever admin saved."""
    row = db.query(CallScript).filter(CallScript.id == DEFAULT_SCRIPT_ID).first()
    if row:
        # Heal a pre-multi-script row that predates the name column.
        if not (row.name or "").strip():
            row.name = "Main Script"
            db.commit()
            db.refresh(row)
        return row
    row = CallScript(
        id=DEFAULT_SCRIPT_ID,
        name="Main Script",
        content=INITIAL_SCRIPT,
        sort_order=0,
        updated_at=_now(),
        updated_by="system",
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


def _list(db) -> list[CallScript]:
    _ensure_seed(db)  # guarantee at least one script exists
    return (
        db.query(CallScript)
        .order_by(CallScript.sort_order, CallScript.name)
        .all()
    )


# ── List / read ───────────────────────────────────────────────────────────
@router.get("/call-scripts")
def list_call_scripts(user: dict = Depends(get_current_user)):
    """Full script library (all roles can read). Ordered for the dropdown."""
    del user
    db = get_db()
    try:
        return {"scripts": [s.to_dict() for s in _list(db)]}
    finally:
        db.close()


@router.get("/call-script")
def get_call_script(user: dict = Depends(get_current_user)):
    """Back-compat single-script read — returns the first (default) script.
    Kept so any older client keeps working; new UI uses /call-scripts."""
    del user
    db = get_db()
    try:
        return _ensure_seed(db).to_dict()
    finally:
        db.close()


# ── Create / update / delete ──────────────────────────────────────────────
@router.post("/call-scripts")
def create_call_script(body: CallScriptCreate, user: dict = Depends(require_admin)):
    name = (body.name or "").strip()
    if not name:
        raise HTTPException(400, "Give the script a name.")
    db = get_db()
    try:
        _ensure_seed(db)
        max_order = db.query(func.max(CallScript.sort_order)).scalar() or 0
        row = CallScript(
            id=str(uuid.uuid4()),
            name=name,
            content=body.content or "",
            sort_order=int(max_order) + 1,
            updated_at=_now(),
            updated_by=user.get("name", ""),
        )
        db.add(row)
        db.commit()
        db.refresh(row)
        return row.to_dict()
    finally:
        db.close()


@router.put("/call-scripts/{script_id}")
def update_call_script_by_id(script_id: str, body: CallScriptUpdate, user: dict = Depends(require_admin)):
    db = get_db()
    try:
        row = db.query(CallScript).filter(CallScript.id == script_id).first()
        if not row:
            raise HTTPException(404, "Script not found")
        if body.name is not None:
            new_name = body.name.strip()
            if not new_name:
                raise HTTPException(400, "Script name can't be empty.")
            row.name = new_name
        if body.content is not None:
            row.content = body.content
        row.updated_at = _now()
        row.updated_by = user.get("name", "")
        db.commit()
        db.refresh(row)
        return row.to_dict()
    finally:
        db.close()


@router.delete("/call-scripts/{script_id}")
def delete_call_script(script_id: str, user: dict = Depends(require_admin)):
    del user
    db = get_db()
    try:
        # Never leave the library empty — the panel needs at least one script.
        if db.query(CallScript).count() <= 1:
            raise HTTPException(400, "Can't delete the only script — add another first.")
        row = db.query(CallScript).filter(CallScript.id == script_id).first()
        if not row:
            raise HTTPException(404, "Script not found")
        db.delete(row)
        db.commit()
        return {"deleted": script_id}
    finally:
        db.close()


# ── Back-compat single-script update (updates the default script) ──────────
@router.put("/call-script")
def update_call_script(body: CallScriptBody, user: dict = Depends(require_admin)):
    db = get_db()
    try:
        row = _ensure_seed(db)
        row.content = body.content or ""
        row.updated_at = _now()
        row.updated_by = user.get("name", "")
        db.commit()
        db.refresh(row)
        return row.to_dict()
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally:
        db.close()


# ── Import: extract text from an uploaded file ─────────────────────────────
def _clean_extracted_text(raw: str) -> str:
    """Tidy raw extracted text into a readable script. Normalizes line endings,
    de-hyphenates words split across line breaks, collapses runs of blank
    lines, and trims trailing whitespace per line — turns a file dump into
    'normal words' the admin can review and save as-is."""
    text = (raw or "").replace("\r\n", "\n").replace("\r", "\n")
    # Join words hyphenated across a line break: "resto-\nration" -> "restoration"
    text = re.sub(r"-\n(\w)", r"\1", text)
    lines = [ln.rstrip() for ln in text.split("\n")]
    out: list[str] = []
    blanks = 0
    for ln in lines:
        if ln.strip():
            blanks = 0
            out.append(ln)
        else:
            blanks += 1
            if blanks <= 1:   # collapse 2+ blank lines into one
                out.append("")
    return "\n".join(out).strip()


def _extract_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        pages = [page.get_text("text") for page in doc]
    finally:
        doc.close()
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    # Pull table cell text too — scripts sometimes live in a 2-column table.
    for table in doc.tables:
        for trow in table.rows:
            cells = [c.text.strip() for c in trow.cells if c.text.strip()]
            if cells:
                parts.append(" — ".join(cells))
    return "\n".join(parts)


async def _extract_upload_text(file: UploadFile) -> str:
    """Read an uploaded PDF / .docx / .txt / .md and return its plain text.
    Raises HTTPException with a friendly message on unsupported/empty/unreadable
    files. Never stores the file — text only."""
    name = (file.filename or "").lower()
    data = await file.read()
    if not data:
        raise HTTPException(400, "That file was empty.")
    try:
        if name.endswith(".pdf"):
            raw = _extract_pdf(data)
            if not raw.strip():
                raise HTTPException(422, "No text found — this looks like a scanned/image PDF, so there's nothing to import.")
        elif name.endswith(".docx"):
            raw = _extract_docx(data)
        elif name.endswith((".txt", ".md")):
            raw = data.decode("utf-8", errors="replace")
        elif name.endswith(".doc"):
            raise HTTPException(400, "Old .doc files aren't supported — open it in Word and 'Save As' .docx (or export a PDF), then upload that.")
        else:
            raise HTTPException(400, "Unsupported file — upload a PDF, Word .docx, or a .txt/.md text file.")
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"Call-script extract failed for {name!r}: {e}")
        raise HTTPException(400, "Couldn't read that file — it may be corrupt or not really that format.")
    text = _clean_extracted_text(raw)
    if not text:
        raise HTTPException(422, "No text found in that file — nothing to import.")
    return text


@router.post("/call-scripts/extract")
async def extract_call_script_file(file: UploadFile = File(...), user: dict = Depends(require_admin)):
    """Pull the text out of an uploaded PDF / Word (.docx) / text file so admin
    can import a script written elsewhere instead of retyping it. Returns the
    extracted text; the caller loads it into the editor to review + save. Does
    NOT save on its own, and never stores the raw file."""
    del user
    text = await _extract_upload_text(file)
    return {"text": text}


@router.post("/call-script/extract-pdf")
async def extract_call_script_pdf(file: UploadFile = File(...), user: dict = Depends(require_admin)):
    """Back-compat alias — now accepts PDF / .docx / text too (name kept so
    older clients keep working)."""
    del user
    text = await _extract_upload_text(file)
    # 'pages' kept in the response shape for the old client; not meaningful for
    # non-PDF sources, so report 1.
    return {"text": text, "pages": 1}
